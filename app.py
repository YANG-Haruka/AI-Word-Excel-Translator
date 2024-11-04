import gradio as gr
import pandas as pd
from docx import Document
from ollama_bot import populate_sum_model, translate_text
from tempfile import NamedTemporaryFile
import os


LANGUAGE_MAP = {
    "日本語": "ja",
    "中文": "zh",
    "English": "en",
}

def translate_excel(file, model, src_lang, dest_lang):
    df = pd.read_excel(file.name, sheet_name=None, engine='openpyxl')
    translated_sheets = {}
    for sheet_name, sheet in df.items():
        translated_sheet = sheet.copy()
        for col in sheet.columns:
            for idx, cell_value in sheet[col].items():
                if isinstance(cell_value, str):
                    translated_text = translate_text(cell_value, model, src_lang, dest_lang)
                    translated_sheet.at[idx, col] = translated_text
        translated_sheets[sheet_name] = translated_sheet
    
    translated_file = NamedTemporaryFile(delete=False, suffix='.xlsx')
    with pd.ExcelWriter(translated_file.name, engine='openpyxl') as writer:
        for sheet_name, translated_sheet in translated_sheets.items():
            translated_sheet.to_excel(writer, sheet_name=sheet_name, index=False)
    return translated_file.name

def translate_word(file, model, src_lang, dest_lang):
    doc = Document(file.name)
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            translated_text = translate_text(paragraph.text, model,src_lang, dest_lang)
            paragraph.text = translated_text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    translated_text = translate_text(cell.text, model, src_lang, dest_lang)
                    cell.text = translated_text
    
    translated_file = NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(translated_file.name)
    return translated_file.name

def translate_file(file, model, src_lang, dest_lang):
    src_lang_code = LANGUAGE_MAP.get(src_lang, "en")
    dest_lang_code = LANGUAGE_MAP.get(dest_lang, "en")

    file_name, file_extension = os.path.splitext(file.name)
    if file_extension == '.xlsx':
        translated_file_path = translate_excel(file, model, src_lang_code, dest_lang_code)
    elif file_extension == '.docx':
        translated_file_path = translate_word(file, model, src_lang_code, dest_lang_code)
    else:
        return "Unsupported file type. Please upload a .docx or .xlsx file."
    
    translated_file_name = f"{file_name}_translated{file_extension}"
    os.rename(translated_file_path, translated_file_name)
    return translated_file_name

available_models = populate_sum_model()

with gr.Blocks() as demo:
    with gr.Row():
        src_lang = gr.Dropdown(["English", "中文", "日本語"], label="Source Language", value="English")
        dest_lang = gr.Dropdown(["English", "中文", "日本語"], label="Target Language", value="English")
        model_choice = gr.Dropdown(available_models, label="Model", value=available_models[0] if available_models else None)

    file_input = gr.File(label="Upload Word/Excel File (.docx or .xlsx)")
    output = gr.File(label="Download Translated File")

    translate_button = gr.Button("Translate")
    translate_button.click(translate_file, inputs=[file_input, model_choice, src_lang, dest_lang], outputs=output)

demo.launch()